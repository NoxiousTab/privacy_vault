import pytest
from core.parser.docx_parser import parse_docx
from pathlib import Path
from docx import Document

TEST_FILE = Path(__file__).parent / "data" / "sample1.docx"

def test_parse_docx_success():
    result = parse_docx(str(TEST_FILE))
    
    expected_content = (
        "This is a sample text file.\n"
        "It contains a few lines to test parsing functionality."
    )

    assert result["filename"] == "sample1.docx"
    assert result["content"] == expected_content
    assert result["word_count"] == len(expected_content.split())

def test_parse_docx_file_not_found():
    with pytest.raises(FileNotFoundError):
        parse_docx("nonexistent.docx")

def test_parse_docx_wrong_extension():
    fake_file = Path(__file__).parent / "data" / "fake.txt"
    fake_file.write_text("fake content")

    with pytest.raises(ValueError):
        parse_docx(str(fake_file))

    fake_file.unlink()

# Optional: Auto-generate the sample file if it doesn't exist
def setup_module(module):
    if not TEST_FILE.exists():
        doc = Document()
        doc.add_paragraph("This is a sample DOCX document.")
        doc.add_paragraph("It is used for unit testing.")
        doc.save(TEST_FILE)
